"""Render resume markdown to ATS-safe PDF (reportlab) and DOCX (python-docx).

Both formats use standard system fonts (Helvetica / Calibri), a real text
layer, and clickable hyperlinks for email / LinkedIn / GitHub — the things
browser print-to-PDF gets wrong for ATS parsers.
"""

import io
import re

# ---------------------------------------------------------------- parsing

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def parse(md: str) -> dict:
    """Parse the fixed resume skeleton into a structure both renderers share."""
    doc = {"name": "", "contact": [], "summary": [], "sections": []}
    current = None
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# ") and not doc["name"]:
            doc["name"] = _plain(line[2:])
            continue
        if line.startswith("## "):
            current = {"title": _plain(line[3:]), "items": []}
            doc["sections"].append(current)
            continue
        if current is None:
            if not doc["contact"] and ("|" in line or "@" in line):
                doc["contact"] = [p.strip() for p in line.split("|") if p.strip()]
            else:
                doc["summary"].append(line)
            continue
        if line.startswith(("- ", "* ")):
            current["items"].append(("bullet", line[2:].strip()))
        elif re.match(r"^\*\*[^*]+\*\*\s*\|", line) and not re.match(r"^\*\*[^*]*:\*\*", line):
            left, _, right = line.partition("|")
            current["items"].append(("role", _plain(left), right.strip()))
        else:
            current["items"].append(("text", line))
    return doc


def _plain(text: str) -> str:
    return BOLD_RE.sub(r"\1", text).strip()


def link_for(token: str):
    """Return (url, display_text) for a contact token; url is None for plain text."""
    t = token.strip().rstrip(",")
    low = t.lower()
    if "@" in t and " " not in t and not any(d in low for d in ("linkedin.", "github.")):
        return ("mailto:" + t, t)
    if any(d in low for d in ("linkedin.com", "github.com", "http://", "https://", "www.")):
        url = t if low.startswith("http") else "https://" + t
        display = re.sub(r"^https?://(www\.)?", "", t)
        return (url, display)
    return (None, t)


# ---------------------------------------------------------------- PDF

def to_pdf(md: str) -> bytes:
    """Render to PDF, auto-shrinking the type until it fits exactly one page."""
    data = b""
    for scale in (1.0, 0.94, 0.88, 0.82, 0.76):
        data, pages = _build_pdf(md, scale)
        if pages <= 1:
            return data
    return data


def _build_pdf(md: str, scale: float):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_RIGHT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    doc = parse(md)
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
        title=f"{doc['name']} - Resume", author=doc["name"],
    )
    usable = LETTER[0] - 1.2 * inch
    z = scale  # type scale factor; < 1.0 compresses everything proportionally

    s_name = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=15.5 * z, leading=18 * z, spaceAfter=1.5 * z)
    s_contact = ParagraphStyle("contact", fontName="Helvetica", fontSize=9 * z, leading=11 * z, spaceAfter=5 * z, textColor=colors.HexColor("#222222"))
    s_body = ParagraphStyle("body", fontName="Helvetica", fontSize=9.5 * z, leading=11.6 * z, spaceAfter=1.5 * z, alignment=TA_JUSTIFY)
    s_h2 = ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=10 * z, leading=12 * z, spaceBefore=6.5 * z, spaceAfter=1 * z)
    s_role = ParagraphStyle("role", fontName="Helvetica-Bold", fontSize=9.5 * z, leading=11.6 * z)
    s_date = ParagraphStyle("date", fontName="Helvetica", fontSize=9 * z, leading=11.6 * z, alignment=TA_RIGHT, textColor=colors.HexColor("#444444"))
    s_bullet = ParagraphStyle("bullet", parent=s_body, leftIndent=12 * z, bulletIndent=3 * z, spaceAfter=1 * z)

    def markup(text: str) -> str:
        out = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", out)

    story = [Paragraph(markup(doc["name"]), s_name)]

    parts = []
    for token in doc["contact"]:
        url, text = link_for(token)
        text = markup(text)
        parts.append(f'<link href="{url}" color="#1a56b0"><u>{text}</u></link>' if url else text)
    if parts:
        story.append(Paragraph("  |  ".join(parts), s_contact))
    for line in doc["summary"]:
        story.append(Paragraph(markup(line), s_body))

    for section in doc["sections"]:
        story.append(Paragraph(section["title"].upper(), s_h2))
        story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#999999"), spaceAfter=3 * z))
        for item in section["items"]:
            if item[0] == "role":
                url, right_text = link_for(item[2])
                right = (
                    f'<link href="{url}" color="#1a56b0"><u>{markup(right_text)}</u></link>'
                    if url else markup(item[2])
                )
                split = 0.58 if len(right_text) > 22 else 0.72
                row = Table(
                    [[Paragraph(markup(item[1]), s_role), Paragraph(right, s_date)]],
                    colWidths=[usable * split, usable * (1 - split)],
                )
                row.setStyle(TableStyle([
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 2.5 * z),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(row)
            elif item[0] == "bullet":
                story.append(Paragraph(markup(item[1]), s_bullet, bulletText="•"))
            else:
                story.append(Paragraph(markup(item[1]), s_body))
        story.append(Spacer(1, 2 * z))

    pdf.build(story)
    return buf.getvalue(), pdf.page


# ---------------------------------------------------------------- DOCX

def to_docx(md: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = parse(md)
    d = Document()

    for section in d.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.6)
    normal = d.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.space_before = Pt(0)

    def add_hyperlink(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color"); color.set(qn("w:val"), "1A56B0"); rpr.append(color)
        under = OxmlElement("w:u"); under.set(qn("w:val"), "single"); rpr.append(under)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        link.append(run)
        paragraph._p.append(link)

    def bottom_border(paragraph):
        ppr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "999999")
        pbdr.append(bottom)
        ppr.append(pbdr)

    def add_rich(paragraph, text):
        """Add text with **bold** segments as runs."""
        for i, chunk in enumerate(BOLD_RE.split(text)):
            if not chunk:
                continue
            run = paragraph.add_run(chunk)
            run.bold = i % 2 == 1

    name_p = d.add_paragraph()
    name_run = name_p.add_run(doc["name"])
    name_run.bold = True
    name_run.font.size = Pt(17)

    contact_p = d.add_paragraph()
    for i, token in enumerate(doc["contact"]):
        if i:
            contact_p.add_run("  |  ")
        url, text = link_for(token)
        if url:
            add_hyperlink(contact_p, url, text)
        else:
            contact_p.add_run(text)
    contact_p.paragraph_format.space_after = Pt(6)

    for line in doc["summary"]:
        p = d.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich(p, line)

    tab_right = Inches(7.3)
    for section in doc["sections"]:
        h = d.add_paragraph()
        run = h.add_run(section["title"].upper())
        run.bold = True
        run.font.size = Pt(11)
        h.paragraph_format.space_before = Pt(8)
        bottom_border(h)
        for item in section["items"]:
            if item[0] == "role":
                p = d.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.tab_stops.add_tab_stop(tab_right, WD_TAB_ALIGNMENT.RIGHT)
                title_run = p.add_run(item[1])
                title_run.bold = True
                p.add_run("\t")
                url, right_text = link_for(item[2])
                if url:
                    add_hyperlink(p, url, right_text)
                else:
                    date_run = p.add_run(item[2])
                    date_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                    date_run.font.size = Pt(10)
            elif item[0] == "bullet":
                p = d.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_rich(p, item[1])
            else:
                p = d.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_rich(p, item[1])

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def letter_to_docx(text: str) -> bytes:
    """Plain business-letter DOCX (Calibri 11, 1in margins)."""
    from docx import Document
    from docx.shared import Inches, Pt

    d = Document()
    for section in d.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
    normal = d.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    for block in text.strip().split("\n\n"):
        d.add_paragraph(block.strip())
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
